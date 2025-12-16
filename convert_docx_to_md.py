from docx import Document
import os

def docx_to_markdown(docx_path, md_path):
    """Convert a DOCX file to Markdown"""
    doc = Document(docx_path)
    
    md_content = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            md_content.append("")
            continue
        
        # Check paragraph style to determine heading level
        style = para.style.name
        
        if 'Heading 1' in style:
            md_content.append(f"# {text}")
        elif 'Heading 2' in style:
            md_content.append(f"## {text}")
        elif 'Heading 3' in style:
            md_content.append(f"### {text}")
        elif 'Heading 4' in style:
            md_content.append(f"#### {text}")
        else:
            # Process inline formatting
            formatted_text = ""
            for run in para.runs:
                run_text = run.text
                if run.bold:
                    run_text = f"**{run_text}**"
                if run.italic:
                    run_text = f"*{run_text}*"
                formatted_text += run_text
            
            md_content.append(formatted_text)
    
    # Handle tables
    for table in doc.tables:
        md_content.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            md_content.append("| " + " | ".join(cells) + " |")
            if i == 0:
                md_content.append("|" + "|".join(["---"] * len(cells)) + "|")
        md_content.append("")
    
    # Write to markdown file
    with open(md_path, 'w', encoding='utf-8') as f:
        f.write("\n".join(md_content))
    
    print(f"Converted: {docx_path}")
    print(f"Saved to: {md_path}")

if __name__ == "__main__":
    docx_file = r"c:\Users\ruobin Yu\Desktop\Special Issue.docx"
    md_file = r"c:\Users\ruobin Yu\.vscode\PoetryAI-6\Manuscript\special_issue\Special_Issue.md"
    
    if os.path.exists(docx_file):
        docx_to_markdown(docx_file, md_file)
    else:
        print(f"Error: File not found - {docx_file}")
